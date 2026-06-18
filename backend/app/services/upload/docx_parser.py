"""
DOCX Parser - Microsoft Word Documents
Extracts text, tables, headings, and metadata
"""
from pathlib import Path
from typing import Dict, List
from docx import Document
from docx.document import Document as DocumentType

from app.utils import logger, ExtractionError


class DOCXParser:
    """Parse Microsoft Word (.docx) documents"""
    
    def extract(self, file_path: str) -> Dict:
        """
        Extract content from DOCX file.
        
        Returns:
            {
                'full_text': str,
                'paragraphs': [{'text': str, 'style': str}],
                'tables': [list of tables],
                'headings': [{'level': int, 'text': str}],
                'metadata': dict,
                'stats': dict
            }
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"DOCX not found: {file_path}")
        
        logger.info(f"Parsing DOCX: {path.name}")
        
        try:
            doc = Document(file_path)
            
            paragraphs = self._extract_paragraphs(doc)
            tables = self._extract_tables(doc)
            headings = self._extract_headings(doc)
            metadata = self._extract_metadata(doc)
            
            # Build full text combining everything
            full_text_parts = []
            
            for para in paragraphs:
                if para['text'].strip():
                    full_text_parts.append(para['text'])
            
            # Add tables as formatted text
            for i, table in enumerate(tables, 1):
                full_text_parts.append(f"\n[Table {i}]")
                full_text_parts.append(self._table_to_text(table))
            
            full_text = "\n\n".join(full_text_parts)
            
            result = {
                'full_text': full_text,
                'paragraphs': paragraphs,
                'tables': tables,
                'headings': headings,
                'metadata': metadata,
                'stats': {
                    'paragraphs_count': len(paragraphs),
                    'tables_count': len(tables),
                    'headings_count': len(headings),
                    'total_characters': len(full_text),
                    'word_count': len(full_text.split())
                }
            }
            
            logger.info(
                f"DOCX extracted: {result['stats']['paragraphs_count']} paragraphs, "
                f"{result['stats']['tables_count']} tables, "
                f"{result['stats']['word_count']} words"
            )
            
            return result
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ExtractionError("DOCX", str(e))
    
    def _extract_paragraphs(self, doc: DocumentType) -> List[Dict]:
        """Extract all paragraphs with style info"""
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            paragraphs.append({
                'text': text,
                'style': para.style.name if para.style else 'Normal',
                'is_heading': para.style.name.startswith('Heading') if para.style else False
            })
        
        return paragraphs
    
    def _extract_tables(self, doc: DocumentType) -> List[List[List[str]]]:
        """Extract all tables as 2D arrays"""
        tables = []
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                if any(row_data):  # Skip empty rows
                    table_data.append(row_data)
            
            if table_data:
                tables.append(table_data)
        
        return tables
    
    def _extract_headings(self, doc: DocumentType) -> List[Dict]:
        """Extract headings with their levels"""
        headings = []
        
        for para in doc.paragraphs:
            if not para.style or not para.style.name.startswith('Heading'):
                continue
            
            text = para.text.strip()
            if not text:
                continue
            
            # Get heading level (Heading 1, Heading 2, etc.)
            try:
                level = int(para.style.name.replace('Heading ', ''))
            except (ValueError, AttributeError):
                level = 1
            
            headings.append({
                'level': level,
                'text': text
            })
        
        return headings
    
    def _extract_metadata(self, doc: DocumentType) -> Dict:
        """Extract document metadata"""
        try:
            props = doc.core_properties
            return {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'keywords': props.keywords or '',
                'created': str(props.created) if props.created else None,
                'modified': str(props.modified) if props.modified else None,
            }
        except Exception:
            return {}
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert table to readable text format"""
        if not table:
            return ""
        
        lines = []
        for row in table:
            # Tab-separated for readability
            line = " | ".join(cell for cell in row)
            lines.append(line)
        
        return "\n".join(lines)


# Global instance
docx_parser = DOCXParser()
